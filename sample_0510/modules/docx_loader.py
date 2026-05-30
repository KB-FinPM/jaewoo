from docx import Document


def extract_docx_section(
    file_path,
    start_title,
    end_title
):
    """
    특정 제목부터 종료 제목까지
    Heading + Paragraph + Table 추출

    Parameters
    ----------
    file_path : str
        DOCX 파일 경로

    start_title : str
        시작 제목

    end_title : str
        종료 제목

    Returns
    -------
    list
        추출 결과 리스트
    """

    doc = Document(file_path)

    is_collecting = False

    results = []

    # =====================================
    # body 순회
    # =====================================

    for element in doc.element.body:

        # =====================================
        # 문단 처리
        # =====================================
        if element.tag.endswith("p"):

            para = next(
                p for p in doc.paragraphs
                if p._element == element
            )

            text = para.text.strip()

            if not text:
                continue

            style = para.style.name

            # ---------------------------------
            # Heading 처리
            # ---------------------------------
            if style.startswith("Heading"):

                # 시작 제목 발견
                if text == start_title:

                    is_collecting = True

                    results.append({
                        "type": "heading",
                        "style": style,
                        "text": text
                    })

                    continue

                # 시작 전이면 무시
                if not is_collecting:
                    continue

                # 종료 제목 발견
                if text == end_title:

                    results.append({
                        "type": "heading",
                        "style": style,
                        "text": text
                    })

                    break

                # 수집 중 Heading 저장
                results.append({
                    "type": "heading",
                    "style": style,
                    "text": text
                })

            # ---------------------------------
            # 일반 본문
            # ---------------------------------
            else:

                if not is_collecting:
                    continue

                results.append({
                    "type": "paragraph",
                    "text": text
                })

        # =====================================
        # 표 처리
        # =====================================
        elif element.tag.endswith("tbl"):

            if not is_collecting:
                continue

            table = next(
                t for t in doc.tables
                if t._element == element
            )

            table_data = []

            for row in table.rows:

                row_data = []

                for cell in row.cells:
                    row_data.append(cell.text.strip())

                table_data.append(row_data)

            results.append({
                "type": "table",
                "data": table_data
            })

    return results