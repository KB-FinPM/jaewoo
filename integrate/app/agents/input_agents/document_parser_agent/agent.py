# EN: Input agent for converting supported uploaded files into structured text.
# KO: 지원되는 업로드 파일을 구조화된 텍스트로 변환하는 Input Agent입니다.

from pathlib import PurePath

from typing import List

from app.core.pm_config import CHUNK_MAX_CHARS, CHUNK_OVERLAP_CHARS
from app.rag.pm_chunking import make_id, semantic_chunk
from app.storage.docx_reader import read_docx_text
from app.agents.input_agents.document_parser_agent.extractor import extract_requirement_atoms
from app.core.pm_logger import log_info, log_step
from app.rag.qdrant_store import QdrantRequirementStore
from app.schemas.pm_artifacts import RequirementAtom
from app.storage.version_manager import (
    build_doc_version_info,
    get_saved_document_info,
    is_version_analyzed,
    load_atoms_cache,
    save_atoms_cache,
)

from app.schemas.io_agent import (
    InputAgentRequest,
    InputAgentResponse,
    InputType,
    NormalizedRequestType,
)


class DocumentParserAgent:
    """Converts external document bytes into structured text for ingestion and local DOCX analysis."""

    def __init__(self, store: QdrantRequirementStore | None = None) -> None:
        self.store = store

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".json", ".log", ".docx"}
    AGENT_NAME = "DocumentParserAgent"

    async def parse(self, request: InputAgentRequest) -> InputAgentResponse:
        if request.input_type != InputType.FILE or not request.files:
            return InputAgentResponse(
                success=False,
                agent_name=self.AGENT_NAME,
                normalized_request_type=NormalizedRequestType.DOCUMENT_INGESTION,
                error="file input is required",
                validation_errors=["file input is required"],
            )

        file_payload = request.files[0]
        extension = PurePath(file_payload.file_name).suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            return InputAgentResponse(
                success=False,
                agent_name=self.AGENT_NAME,
                normalized_request_type=NormalizedRequestType.DOCUMENT_INGESTION,
                error="unsupported file extension",
                validation_errors=["unsupported file extension"],
            )

        text = self._parse_docx_bytes(file_payload.file_bytes) if extension == ".docx" else self._decode_text(file_payload.file_bytes)
        if not text.strip():
            return InputAgentResponse(
                success=False,
                agent_name=self.AGENT_NAME,
                normalized_request_type=NormalizedRequestType.DOCUMENT_INGESTION,
                error="empty parsed text",
                validation_errors=["empty parsed text"],
            )

        return InputAgentResponse(
            agent_name=self.AGENT_NAME,
            normalized_request_type=NormalizedRequestType.DOCUMENT_INGESTION,
            structured_context={
                "text": text,
                "metadata": {
                    "file_name": file_payload.file_name,
                    "extension": extension,
                    "byte_size": len(file_payload.file_bytes),
                    "content_type": file_payload.content_type,
                },
            },
        )

    def _decode_text(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "cp949"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue

        return file_bytes.decode("utf-8", errors="ignore")

    def _parse_docx_bytes(self, file_bytes: bytes) -> str:
        from tempfile import NamedTemporaryFile
        from docx import Document

        with NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
            tmp.write(file_bytes)
            tmp.flush()
            doc = Document(tmp.name)
            parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)

    @staticmethod
    def deduplicate_atoms(atoms: List[RequirementAtom]) -> List[RequirementAtom]:
        seen = set()
        result = []
        for atom in atoms:
            key = (
                atom.category.strip(),
                atom.requirement_name.strip(),
                atom.requirement_type.strip(),
                atom.description.strip()[:120],
            )
            if key in seen:
                continue
            seen.add(key)
            result.append(atom)
        return result

    @staticmethod
    def assign_requirement_ids(atoms: List[RequirementAtom]) -> List[RequirementAtom]:
        for idx, atom in enumerate(atoms, start=1):
            atom.requirement_id = f"REQ-{idx:04d}"
        return atoms

    def analyze(self, docx_path: str, output_dir: str, recreate_collection: bool = False) -> List[RequirementAtom]:
        if not hasattr(self, "store") or self.store is None:
            self.store = QdrantRequirementStore()
        doc_info = build_doc_version_info(docx_path)
        log_step(f"[버전 확인] 문서={doc_info.file_name}, 버전={doc_info.version}, key={doc_info.doc_key}")
        self.store.create_collection(recreate=recreate_collection)

        if is_version_analyzed(output_dir, doc_info):
            log_step("동일 파일로 확인되어 기존 분석 결과를 사용합니다.")
            cached_atoms = load_atoms_cache(output_dir, doc_info)
            if cached_atoms:
                log_info(f"기존 분석 결과 사용: {len(cached_atoms)}건")
                return cached_atoms
            log_info("캐시 파일이 없어 Qdrant에서 기존 분석 결과를 조회합니다.")
            qdrant_atoms = self.store.scroll_atoms_by_doc_key(doc_info.doc_key)
            if qdrant_atoms:
                log_info(f"Qdrant 기존 분석 결과 사용: {len(qdrant_atoms)}건")
                return qdrant_atoms
            log_info("기존 분석 메타데이터는 있으나 데이터가 없어 재분석합니다.")
        else:
            saved_info = get_saved_document_info(output_dir, doc_info)
            if saved_info:
                log_step("파일명은 같지만 생성/수정시각, 파일용량 또는 해시가 달라 새 파일로 판단했습니다. 재분석합니다.")

        if hasattr(self.store, "delete_atoms_by_doc_key"):
            self.store.delete_atoms_by_doc_key(doc_info.doc_key)

        log_step("[1] DOCX 읽기")
        text = read_docx_text(docx_path)

        log_step("[2] Semantic Chunking")
        chunks = semantic_chunk(
            text=text,
            doc_id=make_id(doc_info.doc_key),
            source_file=doc_info.file_name,
            max_chars=CHUNK_MAX_CHARS,
            overlap_chars=CHUNK_OVERLAP_CHARS,
        )
        log_info(f"chunk count: {len(chunks)}")

        all_atoms: List[RequirementAtom] = []
        log_step("[3] 요구사항 Atom 추출 및 Qdrant 저장")
        for idx, chunk in enumerate(chunks, start=1):
            log_info(f"  - chunk {idx}/{len(chunks)}: {chunk.title}, chars={len(chunk.text)}")
            atoms = extract_requirement_atoms(chunk, doc_info)
            for atom in atoms:
                atom.requirement_id = f"REQ-{len(all_atoms) + 1:04d}"
                atom.doc_key = doc_info.doc_key
                atom.doc_version = doc_info.version
            all_atoms.extend(atoms)
            self.store.upsert_atoms(atoms)
            log_info(f"    extracted: {len(atoms)}")

        log_step("[4] 중복 제거")
        all_atoms = self.assign_requirement_ids(self.deduplicate_atoms(all_atoms))
        self.store.upsert_atoms(all_atoms)
        save_atoms_cache(output_dir, doc_info, all_atoms)
        log_info(f"final requirement count: {len(all_atoms)}")
        return all_atoms


document_parser_agent = DocumentParserAgent()
