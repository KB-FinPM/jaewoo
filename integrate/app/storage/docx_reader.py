from docx import Document


def read_docx_text(docx_path: str) -> str:
    doc = Document(docx_path)
    parts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)
