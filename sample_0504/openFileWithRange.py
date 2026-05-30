# python -m pip install \
# langchain \
# langchain-core \
# langchain-community \
# langchain-text-splitters \
# langchain-ollama \
# python-docx \
# chromadb
# 분석 대상 DOCX
# doc = Document("./탬플릿/01.구축요건정의서/KB스타뱅킹전용플레임워크구축.docx")
# 시작 / 종료 제목
# START_TITLE = "1. 프로젝트 구축 내용"
# END_TITLE = "2. 시스템 구성 내역"

from docx import Document
import json

# =========================
# 설정
# =========================

# FILE_PATH = "./탬플릿/01.구축요건정의서/KB스타뱅킹전용플레임워크구축.docx"

# START_TITLE = "1. 프로젝트 구축 내용"
# END_TITLE = "2. 시스템 구성 내역"


# =========================
# 사용자 입력
# =========================
print("입력예시\n- 파일명 : ./탬플릿/01.구축요건정의서/KB스타뱅킹전용플레임워크구축.docx")
print("- 시작제목 : 1. 프로젝트 구축 내용")
print("- 종료제목 : 2. 시스템 구성 내역\n")

FILE_PATH = input("DOCX 파일명을 입력하세요: ").strip()

START_TITLE = input("시작 제목을 입력하세요: ").strip()

END_TITLE = input("종료 제목을 입력하세요: ").strip()



# =========================
# DOCX 로드
# =========================

doc = Document(FILE_PATH)

# =========================
# 상태 변수
# =========================

is_collecting = False

results = []

# =========================
# body 순회
# 문단 + 표 순서 유지
# =========================

for element in doc.element.body:

    # =====================================
    # 문단 처리
    # =====================================
    if element.tag.endswith("p"):

        para = next(
            p for p in doc.paragraphs
            if p._element == element
        )

        text = para.text.strip()

        if not text:
            continue

        style = para.style.name

        # ---------------------------------
        # Heading 처리
        # ---------------------------------
        if style.startswith("Heading"):

            # 시작 제목 발견
            if text == START_TITLE:

                is_collecting = True

                results.append({
                    "type": "heading",
                    "style": style,
                    "text": text
                })

                continue

            # 아직 시작 전이면 무시
            if not is_collecting:
                continue

            # 종료 제목 저장 후 종료
            if text == END_TITLE:

                results.append({
                    "type": "heading",
                    "style": style,
                    "text": text
                })

                break

            # 수집 중 Heading 저장
            results.append({
                "type": "heading",
                "style": style,
                "text": text
            })

        # ---------------------------------
        # 일반 본문
        # ---------------------------------
        else:

            # 시작 전이면 무시
            if not is_collecting:
                continue

            results.append({
                "type": "paragraph",
                "text": text
            })

    # =====================================
    # 표 처리
    # =====================================
    elif element.tag.endswith("tbl"):

        # 시작 전이면 무시
        if not is_collecting:
            continue

        table = next(
            t for t in doc.tables
            if t._element == element
        )

        table_data = []

        for row in table.rows:

            row_data = []

            for cell in row.cells:
                row_data.append(cell.text.strip())

            table_data.append(row_data)

        results.append({
            "type": "table",
            "data": table_data
        })

# =========================
# 결과 출력
# =========================

for item in results:

    print("=" * 80)

    # Heading
    if item["type"] == "heading":

        print(f"[{item['style']}]")
        print(item["text"])

    # Paragraph
    elif item["type"] == "paragraph":

        print("[본문]")
        print(item["text"])

    # Table
    elif item["type"] == "table":

        print("[표]")

        for row in item["data"]:
            print(row)

# =========================
# JSON 저장
# =========================

with open("output.json", "w", encoding="utf-8") as f:
    json.dump(results, f, ensure_ascii=False, indent=2)

print("\n✅ 완료")