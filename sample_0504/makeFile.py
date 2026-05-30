from docx import Document

doc = Document()
doc.add_heading("프로젝트 요약", 0)

table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '항목'
hdr_cells[1].text = '내용'
hdr_cells[2].text = '비고'

data = [
    ("일정", "정상 진행", "-"),
    ("품질", "양호", "-"),
    ("리스크", "낮음", "-")
]

for item, content, note in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = content
    row_cells[2].text = note

doc.save("table_report.docx")