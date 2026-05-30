from docx import Document

doc = Document("word_template.docx")

data = {
    "{{프로젝트_제목}}": "AI PM 시스템 구축",
    "{{프로젝트_배경}}": "모든 사람이 사용할 수 있는 프로젝트 구축",
    "{{프로젝트_내용}}": "쉽게 사용할 수 있는 프로젝트 구축\n정말 쉽게 사용할 수 있을 지는 모르겠음"
}

# 템플릿 내에서 내용 치환
for para in doc.paragraphs:
    for key, value in data.items():
        if key in para.text:
            para.text = para.text.replace(key, value)

# 신규 내용 추가
doc.add_paragraph("새로운 결론 내용 추가")

doc.save("table_report1.docx")